from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable
from uuid import UUID

try:  # pragma: no cover - optional dependency
    from docx import Document
except ImportError:  # pragma: no cover - graceful degradation
    Document = None

try:  # pragma: no cover - optional dependency
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:  # pragma: no cover - graceful degradation
    canvas = None
    letter = (612, 792)

from app.core.config import Settings
from app.models.schemas import ChatResponse, ConversationDetail, ExportArtifact, ExportFormat


class ExportService:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def export_chat_response(
        self,
        conversation_id: UUID,
        response: ChatResponse,
        formats: Iterable[ExportFormat],
    ) -> list[ExportArtifact]:
        if not self.settings.export_enabled:
            return []

        artifacts: list[ExportArtifact] = []
        for export_format in formats:
            path = self._build_output_path(conversation_id, export_format)
            if export_format == "json":
                path.write_text(response.model_dump_json(indent=2), encoding="utf-8")
            elif export_format == "docx":
                self._write_docx(path, response.answer, response.system_notes)
            elif export_format == "pdf":
                self._write_pdf(path, response.answer, response.system_notes)
            artifacts.append(
                ExportArtifact(
                    format=export_format,
                    path=str(path),
                    created_at=datetime.now(timezone.utc),
                )
            )
        return artifacts

    def export_conversation(self, detail: ConversationDetail, export_format: ExportFormat) -> Path:
        path = self._build_output_path(detail.id, export_format)
        if export_format == "json":
            path.write_text(detail.model_dump_json(indent=2), encoding="utf-8")
        elif export_format == "docx":
            if Document is None:
                raise RuntimeError("DOCX export requires python-docx to be installed.")
            document = Document()
            document.add_heading(detail.title or "Conversation export", 0)
            if detail.memory_summary:
                document.add_paragraph(detail.memory_summary)
            for turn in detail.turns:
                document.add_heading(f"User: {turn.query}", level=1)
                document.add_paragraph(turn.answer)
            document.save(path)
        elif export_format == "pdf":
            if canvas is None:
                raise RuntimeError("PDF export requires reportlab to be installed.")
            pdf = canvas.Canvas(str(path), pagesize=letter)
            width, height = letter
            y = height - 60
            pdf.setFont("Helvetica-Bold", 16)
            pdf.drawString(40, y, detail.title or "Conversation export")
            y -= 30
            pdf.setFont("Helvetica", 11)
            if detail.memory_summary:
                for line in self._wrap(detail.memory_summary):
                    pdf.drawString(40, y, line)
                    y -= 16
            for turn in detail.turns:
                y -= 12
                if y < 80:
                    pdf.showPage()
                    y = height - 60
                pdf.setFont("Helvetica-Bold", 12)
                pdf.drawString(40, y, f"User: {turn.query[:90]}")
                y -= 18
                pdf.setFont("Helvetica", 11)
                for line in self._wrap(turn.answer):
                    if y < 80:
                        pdf.showPage()
                        y = height - 60
                        pdf.setFont("Helvetica", 11)
                    pdf.drawString(40, y, line)
                    y -= 16
            pdf.save()
        return path

    def _write_docx(self, path: Path, answer: str, notes: list[str]) -> None:
        if Document is None:
            raise RuntimeError("DOCX export requires python-docx to be installed.")
        document = Document()
        document.add_heading("RAG response export", 0)
        document.add_paragraph(answer)
        if notes:
            document.add_heading("System notes", level=1)
            for note in notes:
                document.add_paragraph(note, style="List Bullet")
        document.save(path)

    def _write_pdf(self, path: Path, answer: str, notes: list[str]) -> None:
        if canvas is None:
            raise RuntimeError("PDF export requires reportlab to be installed.")
        pdf = canvas.Canvas(str(path), pagesize=letter)
        width, height = letter
        y = height - 60
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(40, y, "RAG response export")
        y -= 30
        pdf.setFont("Helvetica", 11)
        for line in self._wrap(answer):
            if y < 80:
                pdf.showPage()
                y = height - 60
                pdf.setFont("Helvetica", 11)
            pdf.drawString(40, y, line)
            y -= 16
        if notes:
            y -= 10
            pdf.setFont("Helvetica-Bold", 12)
            pdf.drawString(40, y, "System notes")
            y -= 18
            pdf.setFont("Helvetica", 11)
            for note in notes:
                for line in self._wrap(f"- {note}"):
                    if y < 80:
                        pdf.showPage()
                        y = height - 60
                        pdf.setFont("Helvetica", 11)
                    pdf.drawString(40, y, line)
                    y -= 16
        pdf.save()

    def _build_output_path(self, conversation_id: UUID, export_format: ExportFormat) -> Path:
        self.settings.export_dir.mkdir(parents=True, exist_ok=True)
        return self.settings.export_dir / f"{conversation_id}.{export_format}"

    @staticmethod
    def _wrap(text: str, width: int = 96) -> list[str]:
        words = text.split()
        if not words:
            return [""]

        lines: list[str] = []
        current = words[0]
        for word in words[1:]:
            tentative = f"{current} {word}"
            if len(tentative) <= width:
                current = tentative
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines
